import os
import subprocess
import logging
import json
from typing import Dict, Tuple, List
from docx import Document
from datetime import datetime
import pandas as pd
import platform
import re  # Added for better text replacement

class ResumeManager:
    def __init__(self, df: pd.DataFrame):
        self.df = df
        self.logger = logging.getLogger(__name__)
        
        # Load personal information from JSON file with improved path handling
        self.personal_info = self.load_personal_info()
        print(f"Personal info loaded: {bool(self.personal_info)}")
        
        self.process_all_resumes()
    
    def load_personal_info(self):
        """Load personal information from JSON file with better path handling."""
        # Try multiple possible paths
        possible_paths = [
            'personal_info.json',  # Current directory
            os.path.join(os.getcwd(), 'personal_info.json'),  # Absolute path to current directory
            os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__))), 'personal_info.json'),  # Project root
        ]
        
        for path in possible_paths:
            try:
                print(f"Trying to load personal info from: {path}")
                with open(path, 'r') as file:
                    content = file.read()
                    # Remove comment line if present
                    if content.startswith('#'):
                        content = '\n'.join(content.split('\n')[1:])
                    
                    data = json.loads(content)
                    print(f"Successfully loaded personal info from {path}")
                    return data
            except FileNotFoundError:
                print(f"File not found: {path}")
                continue
            except json.JSONDecodeError as e:
                print(f"JSON error in {path}: {str(e)}")
                continue
        
        print("WARNING: Could not load personal information. Using defaults.")
        return {
            "name": "Simon Gyimah",
            "contact": {
                "email": "simon.gyimah2@gmail.com",
                "phone": "(315) 603-0026",
                "linkedin": "linkedin.com/in/skgyimah"
            }
        }  # Minimal fallback

    def process_all_resumes(self) -> None:
        for _, row in self.df.iterrows():
            self.create_resume_and_cover(row)

    def create_resume_and_cover(self, row):
        # Get job details
        job_id = row.get('job_id', 'unknown')
        job_role = row.get('job_position_title', '')
        company = row.get('company_name', '')
        location = row.get('location', '')
        job_category = row.get('job_category', 'general')
        job_description = row.get('job_description', '')
        
        # My name - for file naming
        my_name = self.personal_info.get('name', "Simon Gyimah")
        
        print(f"Processing resume and cover letter for: {company} - {job_role} (ID: {job_id})")
        
        # 1. TEMPLATE PATH FIXES
        # Normalize job category to avoid spaces and case issues
        template_category = job_category.lower().replace(" ", "_")
        
        # Try multiple template paths in order of preference
        template_paths = [
            f"templates/resume_{template_category}.docx",  # Preferred format
            f"templates/resume_{template_category.replace('_role', '')}.docx",  # Without 'role'
            "templates/resume_general.docx",  # Generic fallback
            "templates/resume.docx",  # Ultimate fallback
        ]
        
        # Find first available resume template
        resume_path = None
        for path in template_paths:
            if os.path.exists(path):
                resume_path = path
                print(f"Using resume template: {resume_path}")
                break
        
        if not resume_path:
            print("ERROR: No resume template found! Please create a default template at templates/resume.docx")
            return
        
        # Similar approach for cover letter
        cover_paths = [
            f"templates/cover_{template_category}.docx",
            f"templates/cover_{template_category.replace('_role', '')}.docx",
            "templates/cover_general.docx",
            "templates/cover.docx",
        ]
        
        # Find first available cover template
        cover_path = None
        for path in cover_paths:
            if os.path.exists(path):
                cover_path = path
                print(f"Using cover template: {cover_path}")
                break
        
        if not cover_path:
            print("WARNING: No cover letter template found! Skipping cover letter generation.")
            cover_path = None
        
        # 2. OUTPUT PATHS
        # Create appropriate directory name
        company_safe = self.sanitize_filename(company)
        job_role_safe = self.sanitize_filename(job_role)
        today = datetime.now().strftime("%d-%b-%Y")
        
        # Format directory name: CompanyName_JobRole_Date
        output_dir = os.path.join('output', f"{company_safe}_{job_role_safe}_{today}")
        
        # Create directory if it doesn't exist
        try:
            os.makedirs(output_dir, exist_ok=True)
            print(f"Created output directory: {output_dir}")
        except Exception as e:
            print(f"Error creating directory {output_dir}: {str(e)}")
            # Fallback to simpler directory name if there's an error
            output_dir = os.path.join('output', str(job_id))
            os.makedirs(output_dir, exist_ok=True)
            print(f"Using fallback directory: {output_dir}")
        
        # Output paths with descriptive names
        output_docx = os.path.join(output_dir, f"{my_name} Resume.docx")
        output_pdf = os.path.join(output_dir, f"{my_name} Resume.pdf")
        cover_output_docx = os.path.join(output_dir, f"{my_name} Cover Letter.docx")
        cover_output_pdf = os.path.join(output_dir, f"{my_name} Cover Letter.pdf")
        
        print(f"Resume will be saved as: {output_docx}")
        if cover_path:
            print(f"Cover letter will be saved as: {cover_output_docx}")
        
        # 3. PROCESS RESUME
        try:
            resume_doc = Document(resume_path)
            # Get skills as a string
            top_skills = row.get('top_skills', '')
            
            # Create the resume with personal info and job-specific customizations
            self.create_resume(resume_doc, job_role, top_skills, location, output_docx, job_description)
            
        except Exception as e:
            print(f"Error creating resume: {str(e)}")
            import traceback
            traceback.print_exc()
        
        # 4. PROCESS COVER LETTER
        if cover_path:
            try:
                cover_doc = Document(cover_path)
                why_company = row.get('why_this_company', '')
                why_me = row.get('why_me', '')
                
                # Create the cover letter with updated date
                self.create_cover_letter(cover_doc, job_role, company, why_company, why_me, cover_output_docx)
                
            except Exception as e:
                print(f"Error creating cover letter: {str(e)}")
                import traceback
                traceback.print_exc()
    
    @staticmethod
    def sanitize_filename(filename: str) -> str:
        """Remove invalid characters from a filename."""
        # Characters not allowed in Windows filenames: \ / : * ? " < > |
        invalid_chars = ['\\', '/', ':', '*', '?', '"', '<', '>', '|']
        for char in invalid_chars:
            filename = filename.replace(char, '-')
        return filename

    def convert_to_pdf(self, docx_path):
        """Convert a docx file to PDF"""
        try:
            # Get the output PDF path by replacing the extension
            pdf_path = docx_path.replace('.docx', '.pdf')
            
            # Check for LibreOffice existence first
            import subprocess
            import os
            
            # Check if LibreOffice exists in common locations
            libreoffice_paths = [
                'libreoffice',                  # Unix/Linux
                'soffice',                      # Alternative command
                'C:\\Program Files\\LibreOffice\\program\\soffice.exe',  # Windows default
                'C:\\Program Files (x86)\\LibreOffice\\program\\soffice.exe',  # Windows x86
            ]
            
            libreoffice_path = None
            for path in libreoffice_paths:
                try:
                    subprocess.run([path, '--version'], capture_output=True, check=False)
                    libreoffice_path = path
                    break
                except (subprocess.SubprocessError, FileNotFoundError):
                    continue
            
            if libreoffice_path:
                # Convert using LibreOffice
                subprocess.run([
                    libreoffice_path,
                    '--headless',
                    '--convert-to', 'pdf',
                    '--outdir', os.path.dirname(pdf_path),
                    docx_path
                ], check=True)
                print(f"PDF created at: {pdf_path}")
            else:
                print("LibreOffice not found. Skipping PDF conversion.")
                
        except Exception as e:
            print(f"Error converting to PDF: {str(e)}")

    def create_cover_letter(self, cover_doc, job_role, company, why_company, why_me, output_path):
        """Create a cover letter document with improved placeholder replacement"""
        try:
            # Get current date
            current_date = datetime.now().strftime("%B %d, %Y")
            
            # Dictionary of replacements to make
            replacements = {
                "[Current Date]": current_date,
                "{{DATE}}": current_date,
                "[DATE]": current_date,
                
                "[JOB_ROLE]": job_role,
                "{{JOB_ROLE}}": job_role,
                
                "[Company Name]": company,
                "{{COMPANY}}": company,
                "[COMPANY]": company,
                "[COMPANY_NAME]": company,
                
                "[why_this_company]": why_company,
                "{{WHY_COMPANY}}": why_company,
                "[WHY_COMPANY]": why_company,
                
                "[why_me]": why_me,
                "{{WHY_ME}}": why_me,
                "[WHY_ME]": why_me,
                
                "[NAME]": self.personal_info.get('name', "Simon Gyimah"),
                "{{NAME}}": self.personal_info.get('name', "Simon Gyimah"),
                
                "[City, State ZIP]": "Stevens Point, WI",  # Adding common placeholders
                "[Company Address]": "Company Address"
            }
            
            # Process paragraphs
            for paragraph in cover_doc.paragraphs:
                original_text = paragraph.text
                modified_text = original_text
                
                # Apply all replacements
                for placeholder, value in replacements.items():
                    if placeholder in modified_text:
                        print(f"Replacing '{placeholder}' with '{value}' in paragraph")
                        modified_text = modified_text.replace(placeholder, value)
                
                # Only reassign if text changed (more efficient)
                if modified_text != original_text:
                    # Handle text runs to preserve formatting
                    for run in paragraph.runs:
                        for placeholder, value in replacements.items():
                            if placeholder in run.text:
                                run.text = run.text.replace(placeholder, value)
            
            # Process tables
            for table in cover_doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            original_text = paragraph.text
                            modified_text = original_text
                            
                            # Apply all replacements
                            for placeholder, value in replacements.items():
                                if placeholder in modified_text:
                                    print(f"Replacing '{placeholder}' with '{value}' in table cell")
                                    modified_text = modified_text.replace(placeholder, value)
                            
                            # Only reassign if text changed
                            if modified_text != original_text:
                                # Handle runs for formatting
                                for run in paragraph.runs:
                                    for placeholder, value in replacements.items():
                                        if placeholder in run.text:
                                            run.text = run.text.replace(placeholder, value)
            
            # Save the document
            cover_doc.save(output_path)
            print(f"Cover letter created at: {output_path}")
            
            # Convert to PDF
            self.convert_to_pdf(output_path)
            
        except Exception as e:
            print(f"Error creating cover letter: {str(e)}")
            import traceback
            traceback.print_exc()

    def create_resume(self, resume_doc, job_role, top_skills, location, output_path, job_description=''):
        """Create a customized resume using personal information and job details with improved replacement"""
        try:
            # Your name from personal info
            my_name = self.personal_info.get('name', "Simon Gyimah")
            
            # Get contact info
            contact = self.personal_info.get('contact', {})
            email = contact.get('email', '')
            phone = contact.get('phone', '')
            linkedin = contact.get('linkedin', '')
            
            # Extract relevant skills for this job
            job_keywords = self.extract_keywords(job_description)
            relevant_skills = self.get_relevant_skills(job_keywords)
            
            # Get education info
            education = self.format_education(self.personal_info.get('education', []))
            
            # Get experience info - prioritize relevant experiences
            experience = self.format_experience(self.personal_info.get('experience', []), job_keywords)
            
            # Generate professional summary
            summary = self.generate_summary(job_role, top_skills, job_keywords)
            
            # Dictionary of replacements to make
            replacements = {
                "[NAME]": my_name,
                "{{NAME}}": my_name,
                
                "[JOB_ROLE]": job_role,
                "{{JOB_ROLE}}": job_role,
                
                "[SKILLS]": relevant_skills,
                "{{TOP_SKILLS}}": relevant_skills,
                "{{SKILLS}}": relevant_skills,
                
                "[LOCATION]": location,
                "{{LOCATION}}": location,
                
                "[SUMMARY]": summary,
                "{{SUMMARY}}": summary,
                
                "[EDUCATION]": education,
                "{{EDUCATION}}": education,
                
                "[EXPERIENCE]": experience,
                "{{EXPERIENCE}}": experience,
                
                "[EMAIL]": email,
                "{{EMAIL}}": email,
                
                "[PHONE]": phone,
                "{{PHONE}}": phone,
                
                "[LINKEDIN]": linkedin,
                "{{LINKEDIN}}": linkedin
            }
            
            # Add generic name replacements
            for name_to_replace in ["John Doe", "Jane Doe", "John Smith", "Jane Smith"]:
                replacements[name_to_replace] = my_name
            
            # Debug: Print all text elements to help diagnose template issues
            print("\nDEBUG - Resume document contents:")
            for i, para in enumerate(resume_doc.paragraphs):
                print(f"Paragraph {i}: {para.text[:50]}..." if len(para.text) > 50 else f"Paragraph {i}: {para.text}")
                # Look for placeholders
                for placeholder in replacements.keys():
                    if placeholder in para.text:
                        print(f"  Found placeholder: {placeholder}")
            
            # Apply replacements to paragraphs
            for paragraph in resume_doc.paragraphs:
                original_text = paragraph.text
                modified_text = original_text
                
                # Apply all replacements
                for placeholder, value in replacements.items():
                    if placeholder in modified_text:
                        print(f"Replacing '{placeholder}' with '{value}' in paragraph")
                        modified_text = modified_text.replace(placeholder, value)
                
                # Only update if changes were made
                if modified_text != original_text:
                    # Handle runs to preserve formatting
                    runs = list(paragraph.runs)
                    for i, run in enumerate(runs):
                        for placeholder, value in replacements.items():
                            if placeholder in run.text:
                                run.text = run.text.replace(placeholder, value)
            
            # Process tables
            for table in resume_doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            original_text = paragraph.text
                            modified_text = original_text
                            
                            # Apply all replacements
                            for placeholder, value in replacements.items():
                                if placeholder in modified_text:
                                    print(f"Replacing '{placeholder}' with '{value}' in table cell")
                                    modified_text = modified_text.replace(placeholder, value)
                            
                            # Only update if changes were made
                            if modified_text != original_text:
                                # Handle runs for formatting
                                for run in paragraph.runs:
                                    for placeholder, value in replacements.items():
                                        if placeholder in run.text:
                                            run.text = run.text.replace(placeholder, value)
            
            # Save document
            resume_doc.save(output_path)
            print(f"Resume created at: {output_path}")
            
            # Try to convert to PDF
            self.convert_to_pdf(output_path)
            
        except Exception as e:
            print(f"Error in create_resume: {str(e)}")
            import traceback
            traceback.print_exc()

    def extract_keywords(self, job_description):
        """Extract keywords from job description"""
        if not job_description:
            return []
            
        # Convert to lowercase and split into words
        words = job_description.lower().split()
        
        # Filter common words, keep only relevant ones
        common_words = {'the', 'and', 'a', 'to', 'of', 'in', 'for', 'with', 'on', 'at', 'from', 'by'}
        keywords = [word for word in words if word not in common_words and len(word) > 3]
        
        # Return unique keywords
        return list(set(keywords))
    
    def get_relevant_skills(self, job_keywords):
        """Get skills relevant to the job keywords"""
        if not job_keywords:
            # Return all skills if no keywords
            all_skills = []
            for skill_type in self.personal_info.get('skills', {}).values():
                all_skills.extend(skill_type)
            return ', '.join(all_skills[:10])  # Limit to 10 skills
            
        # Get all skills
        all_skills = []
        skill_scores = {}
        
        for skill_type, skills in self.personal_info.get('skills', {}).items():
            for skill in skills:
                all_skills.append(skill)
                # Score each skill based on keywords match
                skill_lower = skill.lower()
                score = 0
                for keyword in job_keywords:
                    if keyword in skill_lower:
                        score += 10  # Direct match
                skill_scores[skill] = score
        
        # Add default score for all skills to ensure some are selected
        for skill in all_skills:
            if skill not in skill_scores:
                skill_scores[skill] = 1
        
        # Sort skills by score
        sorted_skills = sorted(all_skills, key=lambda s: skill_scores.get(s, 0), reverse=True)
        
        # Return top skills (limit to 8)
        return ', '.join(sorted_skills[:8])
    
    def format_education(self, education_list):
        """Format education for resume"""
        if not education_list:
            return ""
            
        result = []
        for edu in education_list:
            degree = edu.get('degree', '')
            school = edu.get('school', '')
            year = edu.get('year', '')
            gpa = edu.get('gpa', '')
            
            edu_line = f"{degree}, {school}, {year}"
            if gpa:
                edu_line += f" (GPA: {gpa})"
                
            result.append(edu_line)
            
        return '\n'.join(result)
    
    def format_experience(self, experience_list, job_keywords):
        """Format experiences, prioritizing those relevant to job keywords"""
        if not experience_list:
            return ""
            
        # Score each experience based on relevance to job
        scored_experiences = []
        for exp in experience_list:
            title = exp.get('title', '')
            company = exp.get('company', '')
            dates = exp.get('dates', '')
            location = exp.get('location', '')
            bullets = exp.get('bullets', [])
            skills = exp.get('skills', [])
            
            # Calculate relevance score
            score = 0
            
            # Check title relevance
            for keyword in job_keywords:
                if keyword in title.lower():
                    score += 5
                if keyword in company.lower():
                    score += 3
                    
            # Check bullet points relevance
            for bullet in bullets:
                for keyword in job_keywords:
                    if keyword in bullet.lower():
                        score += 2
                        
            # Check skills relevance
            for skill in skills:
                for keyword in job_keywords:
                    if keyword in skill.lower():
                        score += 4
            
            # Format the experience
            exp_text = f"{title} at {company}, {dates}\n"
            
            # Add bullet points
            for bullet in bullets:
                exp_text += f"• {bullet}\n"
                
            scored_experiences.append((score, exp_text))
            
        # Sort by relevance score (highest first)
        scored_experiences.sort(reverse=True, key=lambda x: x[0])
        
        # Return formatted experiences
        return '\n'.join(exp[1] for exp in scored_experiences)
    
    def generate_summary(self, job_role, top_skills, job_keywords):
        """Generate a professional summary based on the job"""
        # Get highest education
        education = self.personal_info.get('education', [])
        highest_degree = "professional"
        if education:
            highest_degree = education[0].get('degree', 'professional')
            
        # Count years of experience
        experience = self.personal_info.get('experience', [])
        years_experience = len(experience)
        
        # Generate summary with keyword emphasis
        summary = f"{highest_degree} with {years_experience}+ years of experience"
        
        # Identify key skills for this role
        skill_emphasis = self.get_relevant_skills(job_keywords).split(', ')[:3]
        
        if skill_emphasis:
            summary += f" specializing in {', '.join(skill_emphasis)}"
            
        summary += f". Seeking a {job_role} position to leverage analytical abilities and technical expertise in driving data-informed decisions and delivering actionable insights."
        
        return summary

    @staticmethod
    def save_to_pdf(output_path: str, docx_path: str) -> None:
        """Convert docx to PDF using available tools based on platform."""
        try:
            if platform.system() == "Darwin":  # macOS
                libreoffice_path = "/Applications/LibreOffice.app/Contents/MacOS/soffice"
                subprocess.run([
                    libreoffice_path,
                    "--headless",
                    "--convert-to",
                    "pdf",
                    "--outdir",
                    output_path,
                    docx_path
                ], check=True)
            elif platform.system() == "Windows":
                # Windows - try to find LibreOffice or skip PDF conversion
                libreoffice_paths = [
                    r"C:\Program Files\LibreOffice\program\soffice.exe",
                    r"C:\Program Files (x86)\LibreOffice\program\soffice.exe"
                ]
                
                for path in libreoffice_paths:
                    if os.path.exists(path):
                        subprocess.run([
                            path,
                            "--headless",
                            "--convert-to",
                            "pdf",
                            "--outdir",
                            output_path,
                            docx_path
                        ], check=True)
                        return
                
                print("LibreOffice not found. Skipping PDF conversion.")
            else:
                # Linux or other systems
                try:
                    subprocess.run([
                        "libreoffice",
                        "--headless",
                        "--convert-to",
                        "pdf",
                        "--outdir",
                        output_path,
                        docx_path
                    ], check=True)
                except:
                    print("Failed to convert to PDF. LibreOffice may not be installed.")
        except subprocess.CalledProcessError as e:
            print(f"Error converting to PDF: {e}")

if __name__ == "__main__":
    # Add any test or example usage here
    pass