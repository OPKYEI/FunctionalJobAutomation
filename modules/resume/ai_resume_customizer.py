#!/usr/bin/env python3
"""
AI-Powered Resume Customizer

This script generates highly tailored resumes based on job descriptions scraped from LinkedIn.
It uses OpenAI's GPT model to naturally reframe your work experience to highlight relevant skills
while maintaining authenticity. It also generates a concise professional summary and formats
the resume with proper margins and date alignment.

Usage:
    python -m modules.resume.ai_resume_customizer --job_id <job_id> --output <output_filename.pdf>

Requirements:
    - python-docx
    - docx2pdf
    - pandas
"""

import os
import re
import json
import argparse
import pandas as pd
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx2pdf import convert
import logging
import sys

# Add project root to path for importing modules
sys.path.append(os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))))

# Import OpenAI connection from existing modules
from modules.ai.openaiConnections import ai_create_openai_client, ai_close_openai_client

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    datefmt='%Y-%m-%d %H:%M:%S'
)
logger = logging.getLogger(__name__)

# Paths to data files
RESUME_TEMPLATE = "templates/resume_template.docx"
BASE_RESUME_DATA = "data/resume_data.json"
JOB_APPLICATIONS_CSV = "all excels/all_applied_applications_history.csv"
OUTPUT_DIR = "all resumes/customized"

# Create output directory if it doesn't exist
os.makedirs(OUTPUT_DIR, exist_ok=True)

class AIResumeCustomizer:
    def __init__(self, base_resume_path=BASE_RESUME_DATA, template_path=RESUME_TEMPLATE):
        """Initialize the resume customizer with base resume data and template."""
        # Load base resume data
        with open(base_resume_path, 'r') as f:
            self.resume_data = json.load(f)
            
        self.template_path = template_path
        
        # Initialize OpenAI client
        self.ai_client = ai_create_openai_client()
        if not self.ai_client:
            logger.error("Failed to create OpenAI client. Please check your API key.")
            raise ValueError("OpenAI client initialization failed")
        
        # Initialize job-specific customized data
        self.customized_data = None
        
        logger.info("AI-Powered Resume Customizer initialized")
    
    def load_job_description(self, job_id):
        """Load job description from CSV based on job ID."""
        try:
            # Load the applications CSV
            df = pd.read_csv(JOB_APPLICATIONS_CSV)
            
            # Find the job by ID
            job_data = df[df['Job ID'] == job_id].iloc[0]
            
            # Extract relevant job information
            job_info = {
                'job_id': job_id,
                'title': job_data['Title'],
                'company': job_data['Company'],
                'location': job_data['Work Location'],
                'work_style': job_data['Work Style'],
                'description': job_data['About Job'],
                'skills_required': job_data['Skills required']
            }
            
            logger.info(f"Loaded job description for ID {job_id}")
            return job_info
            
        except Exception as e:
            logger.error(f"Error loading job description: {e}")
            raise
    
    def extract_job_requirements(self, job_description):
        """Extract job requirements with OpenAI without requiring JSON format."""
        try:
            system_prompt = """You are an expert at analyzing job descriptions to extract key requirements.
            Extract the following from the job description:
            1. Required technical skills
            2. Soft skills
            3. Job responsibilities
            4. Education requirements
            5. Location requirements (if any)
            6. Experience level required (in years if specified)
            7. Industry/domain focus

            Format your response in this structured way:
            TECHNICAL SKILLS: Skill 1, Skill 2, Skill 3, ...
            SOFT SKILLS: Skill 1, Skill 2, Skill 3, ...
            RESPONSIBILITIES: Responsibility 1; Responsibility 2; Responsibility 3; ...
            EDUCATION: Requirement 1, Requirement 2, ...
            LOCATION: Location (or 'None specified')
            YEARS_EXPERIENCE: Number (or 'Not specified')
            DOMAIN: Industry/domain (or 'Not specified')
            """

            user_prompt = f"Analyze this job description and extract the key requirements:\n\n{job_description}"
            
            response = self.ai_client.chat.completions.create(
                model="gpt-3.5-turbo",  # Using 3.5 for speed and compatibility
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt}
                ]
            )
            
            # Parse the text-based response
            response_text = response.choices[0].message.content
            requirements = self._parse_structured_response(response_text)
            
            logger.info(f"Extracted {len(requirements.get('technical_skills', []))} technical skills, {len(requirements.get('responsibilities', []))} responsibilities")
            
            return requirements
            
        except Exception as e:
            logger.error(f"Error extracting job requirements with AI: {e}")
            # Fallback to basic extraction if AI fails
            return self._basic_extract_requirements(job_description)
    def _parse_structured_response(self, response_text):
        """Parse the structured response from the AI into a dictionary."""
        requirements = {
            "technical_skills": [],
            "soft_skills": [],
            "responsibilities": [],
            "education": [],
            "location_required": None,
            "years_experience": None,
            "domain_focus": None
        }
        
        # Split by lines and process each line
        lines = response_text.strip().split('\n')
        current_section = None
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            # Check for section headers
            if line.upper().startswith("TECHNICAL SKILLS:"):
                current_section = "technical_skills"
                content = line[line.index(':')+1:].strip()
            elif line.upper().startswith("SOFT SKILLS:"):
                current_section = "soft_skills"
                content = line[line.index(':')+1:].strip()
            elif line.upper().startswith("RESPONSIBILITIES:"):
                current_section = "responsibilities"
                content = line[line.index(':')+1:].strip()
            elif line.upper().startswith("EDUCATION:"):
                current_section = "education"
                content = line[line.index(':')+1:].strip()
            elif line.upper().startswith("LOCATION:"):
                location = line[line.index(':')+1:].strip()
                if location.lower() not in ["none", "none specified", "not specified"]:
                    requirements["location_required"] = location
                current_section = None
                continue
            elif line.upper().startswith("YEARS_EXPERIENCE:"):
                experience = line[line.index(':')+1:].strip()
                if experience.lower() not in ["none", "none specified", "not specified"]:
                    try:
                        requirements["years_experience"] = int(experience.split()[0])
                    except:
                        requirements["years_experience"] = experience
                current_section = None
                continue
            elif line.upper().startswith("DOMAIN:"):
                domain = line[line.index(':')+1:].strip()
                if domain.lower() not in ["none", "none specified", "not specified"]:
                    requirements["domain_focus"] = domain
                current_section = None
                continue
            elif ":" in line:  # Handle any other sections we might have missed
                potential_section = line[:line.index(':')].lower().strip()
                if potential_section in requirements:
                    current_section = potential_section
                    content = line[line.index(':')+1:].strip()
                else:
                    content = line
            else:
                content = line
                
            # Process content if we're in a valid section
            if current_section and current_section in requirements:
                if current_section == "responsibilities":
                    # Split responsibilities by semicolons
                    items = [item.strip() for item in content.split(';') if item.strip()]
                    requirements[current_section].extend(items)
                else:
                    # Split other items by commas
                    items = [item.strip() for item in content.split(',') if item.strip()]
                    requirements[current_section].extend(items)
        
        return requirements
    def _basic_extract_requirements(self, job_description):
        """Basic extraction as fallback if AI extraction fails."""
        # Extract location requirements
        location_required = None
        location_patterns = [
            r"candidate must (?:live|be located) in (?:the )?([\w\s,]+) area",
            r"location(?:ed)? in (?:the )?([\w\s,]+) area",
            r"based in (?:the )?([\w\s,]+) area"
        ]
        
        for pattern in location_patterns:
            location_match = re.search(pattern, job_description, re.IGNORECASE)
            if location_match:
                location_required = location_match.group(1).strip()
                break
        
        # Extract basic skills
        skills = []
        skill_patterns = [r'experience (?:with|in) ([\w\s,]+)', r'knowledge of ([\w\s,]+)']
        for pattern in skill_patterns:
            matches = re.findall(pattern, job_description.lower())
            for match in matches:
                skills.extend([s.strip() for s in match.split(',')])
        
        # Return minimal requirements
        return {
            "technical_skills": list(set(skills)),
            "soft_skills": [],
            "responsibilities": [],
            "education": [],
            "location_required": location_required,
            "years_experience": None,
            "domain_focus": None
        }
    
    def generate_professional_summary(self, job_requirements, job_info):
        """Generate a concise professional summary tailored to the job."""
        try:
            personal_info = self.resume_data["personal_info"]
            experiences = self.resume_data["work_experience"]
            skills = self.resume_data["skills"]
            
            # Construct prompt for AI
            prompt = f"""
            Task: Create a concise professional summary (3-4 sentences max) for my resume, tailored to this job.
            
            About me:
            - Name: {personal_info.get('first_name', '')} {personal_info.get('last_name', '')}
            - Current position: {experiences[0].get('title', '') if experiences else 'N/A'} at {experiences[0].get('company', '') if experiences else 'N/A'}
            - Years of experience: {len(experiences)} positions spanning approximately {sum(range(1, len(experiences) + 1))} years
            - Top skills: {', '.join(skills[:5])}
            
            Job I'm applying for:
            - Title: {job_info.get('title', '')}
            - Company: {job_info.get('company', '')}
            - Key job requirements: {', '.join(job_requirements.get('technical_skills', [])[:5])}
            - Key responsibilities: {'; '.join(job_requirements.get('responsibilities', [])[:2])}
            
            Write a concise professional summary (3-4 sentences) that:
            1. Highlights my relevant experience and skills
            2. Shows how my background directly relates to the job requirements
            3. Includes specific numbers/metrics where possible
            4. Positions me as the ideal candidate for this specific role at {job_info.get('company', '')}
            5. Uses active voice and first-person perspective
            
            Professional Summary:
            """
            
            response = self.ai_client.chat.completions.create(
                model="gpt-4",
                messages=[
                    {"role": "user", "content": prompt}
                ]
            )
            
            # Parse and clean the response
            summary = response.choices[0].message.content.strip()
            
            # Remove any titles/headers that may have been added
            summary = re.sub(r'^(Professional Summary:?\s*|Summary:?\s*)', '', summary, flags=re.IGNORECASE)
            
            logger.info("Generated professional summary")
            return summary
            
        except Exception as e:
            logger.error(f"Error generating professional summary with AI: {e}")
            # Return a basic summary if AI fails
            return f"Experienced professional with skills in {', '.join(self.resume_data['skills'][:3])} seeking the {job_info.get('title', 'advertised')} position at {job_info.get('company', '')}."
    
    def customize_resume_data(self, job_requirements, job_info):
        """Create a customized version of resume data tailored to the job requirements."""
        # Create a copy of the original resume data
        customized = self.resume_data.copy()
        
        # Generate professional summary
        professional_summary = self.generate_professional_summary(job_requirements, job_info)
        customized["professional_summary"] = professional_summary
        
        # Adjust personal info for location if needed
        self._customize_location(customized, job_requirements)
        
        # Customize skills section - reorder to prioritize matching skills
        self._customize_skills(customized, job_requirements)
        
        # Customize work experience - use AI to reframe bullet points
        self._customize_work_experience_with_ai(customized, job_requirements, job_info)
        
        # Customize projects section - select most relevant projects
        self._customize_projects(customized, job_requirements)
        
        # Set the customized data
        self.customized_data = customized
        
        return customized
    
    def _customize_location(self, customized_data, job_requirements):
        """Adjust location information if the job has specific location requirements."""
        # Only modify if the job has a specific location requirement
        if job_requirements.get("location_required"):
            # Keep original location data for reference
            original_location = customized_data["personal_info"].get("location", "")
            
            # For the presentation, adjust to show availability in the required location
            required_location = job_requirements["location_required"]
            
            # If the locations are different and it's a relocation scenario
            if required_location.lower() not in original_location.lower():
                # We'll add a note about relocation in the customized data
                customized_data["relocation_note"] = required_location
            
            logger.info(f"Adjusted location information for job requirement: {required_location}")
    
    def _customize_skills(self, customized_data, job_requirements):
        """Reorder skills to prioritize those matching the job requirements."""
        # Get all skills from resume data
        all_skills = customized_data.get("skills", [])
        if not all_skills:
            return
        
        # Use AI to prioritize skills
        try:
            skills_prompt = f"""
            Task: Prioritize these skills based on relevance to the job requirements.
            
            My skills:
            {", ".join(all_skills)}
            
            Job requirements:
            Technical skills: {", ".join(job_requirements.get("technical_skills", []))}
            Responsibilities: {"; ".join(job_requirements.get("responsibilities", [])[:3])}
            
            Return a comma-separated list of my skills, prioritized by relevance to this job.
            Only include skills from my original list, don't add new ones.
            """
            
            response = self.ai_client.chat.completions.create(
                model="gpt-3.5-turbo",  # Using faster model for this simple task
                messages=[
                    {"role": "user", "content": skills_prompt}
                ]
            )
            
            # Parse the response
            skills_text = response.choices[0].message.content
            prioritized_skills = [s.strip() for s in skills_text.split(',')]
            
            # Verify all skills are in the original list
            valid_skills = [s for s in prioritized_skills if s in all_skills]
            
            # Add any missing skills to the end
            missing_skills = [s for s in all_skills if s not in valid_skills]
            valid_skills.extend(missing_skills)
            
            # Update skills in the customized data
            customized_data["skills"] = valid_skills
            
            logger.info(f"Reordered skills list to prioritize job-relevant skills")
            
        except Exception as e:
            logger.error(f"Error customizing skills with AI: {e}")
            # Fallback to keeping original skills if AI fails
    
    def _customize_work_experience_with_ai(self, customized_data, job_requirements, job_info):
        """Use AI to reframe work experience bullet points to match job requirements."""
        # Get responsibilities from job requirements
        responsibilities = job_requirements.get("responsibilities", [])
        if not responsibilities:
            logger.warning("No job responsibilities found to customize experience")
            return
                
        # For each work experience entry
        for exp_idx, experience in enumerate(customized_data.get("work_experience", [])):
            if "bullet_points" not in experience or not experience["bullet_points"]:
                continue
                    
            # Get original bullet points
            original_bullets = experience["bullet_points"]
            company_name = experience["company"]
                
            # Construct prompt for the AI
            prompt = f"""
            Task: Rewrite my work experience bullet points to better match a job description while maintaining authenticity.
            
            My original work experience at {company_name}:
            {"".join(f"• {bullet}\n" for bullet in original_bullets)}
            
            Job responsibilities I want to highlight:
            {"".join(f"• {resp}\n" for resp in responsibilities[:len(original_bullets)])}
            
            Rules:
            1. Keep my actual metrics and achievements (percentages, numbers)
            2. Maintain {company_name} as the workplace
            3. Make each bullet point address one of the job responsibilities
            4. Use strong action verbs at the beginning of each bullet
            5. Ensure complete, grammatically correct sentences
            6. Create exactly {len(original_bullets)} bullet points
            7. Do not truncate or abbreviate text
            8. Do not include bullet point markers (•) in your response
            9. DO NOT use quotation marks around the bullet points
            10. Format each bullet as: [Action verb] [responsibility-focused accomplishment], resulting in [original metric] at {company_name}
            
            Rewritten bullet points (without bullet markers or quotation marks):
            """
            
            try:
                response = self.ai_client.chat.completions.create(
                    model="gpt-3.5-turbo",  # Using 3.5 for speed and compatibility
                    messages=[
                        {"role": "user", "content": prompt}
                    ]
                )
                
                # Parse the response and extract bullet points
                response_text = response.choices[0].message.content
                new_bullets = []
                
                # Clean up the response and extract bullet points
                for line in response_text.split('\n'):
                    line = line.strip()
                    # Skip empty lines or lines with bullet markers 
                    if not line or line.startswith('•'):
                        continue
                    # Clean up any numbering
                    if re.match(r'^\d+\.\s', line):
                        line = re.sub(r'^\d+\.\s', '', line)
                    # Remove any quotation marks
                    line = line.strip('"\'')
                    # Add to bullet points
                    if line:
                        new_bullets.append(line)
                
                # Ensure we have the right number of bullet points
                if len(new_bullets) > len(original_bullets):
                    new_bullets = new_bullets[:len(original_bullets)]
                
                # If we got fewer bullet points, keep some original ones
                while len(new_bullets) < len(original_bullets):
                    new_bullets.append(original_bullets[len(new_bullets)])
                
                # Update bullet points in the customized data
                experience["bullet_points"] = new_bullets
                customized_data["work_experience"][exp_idx] = experience
                
                logger.info(f"Successfully reframed {len(new_bullets)} bullet points for {company_name}")
                
            except Exception as e:
                logger.error(f"Error customizing experience with AI for {company_name}: {e}")
                # Keep original bullets if AI fails
                logger.info(f"Keeping original bullet points for {company_name}")
    
    def _customize_projects(self, customized_data, job_requirements):
        """Select most relevant projects based on job requirements."""
        # Get all projects from resume data
        all_projects = customized_data.get("projects", [])
        if not all_projects:
            return
            
        try:
            # Create a prompt for the AI to rank projects
            projects_text = ""
            for i, project in enumerate(all_projects):
                tech = project.get("technologies", "")
                desc = project.get("description", "")
                projects_text += f"Project {i+1}: {project.get('title', '')}\n"
                projects_text += f"Description: {desc}\n"
                projects_text += f"Technologies: {tech}\n\n"
            
            prompt = f"""
            Task: Rank these projects based on relevance to the job requirements.
            
            Job requirements:
            Technical skills: {", ".join(job_requirements.get("technical_skills", []))}
            Responsibilities: {"; ".join(job_requirements.get("responsibilities", [])[:3])}
            
            My projects:
            {projects_text}
            
            Return a comma-separated list of project numbers (1, 2, 3, etc.) ranked from most relevant to least relevant.
            Only include the project numbers, no other text.
            """
            
            response = self.ai_client.chat.completions.create(
                model="gpt-3.5-turbo",  # Using faster model for simple ranking
                messages=[
                    {"role": "user", "content": prompt}
                ]
            )
            
            # Parse the response
            ranking_text = response.choices[0].message.content.strip()
            # Extract numbers from the response
            project_ranks = []
            for num in re.findall(r'\d+', ranking_text):
                idx = int(num) - 1  # Convert to 0-based index
                if 0 <= idx < len(all_projects) and idx not in project_ranks:
                    project_ranks.append(idx)
            
            # Add any missing indexes to the end
            for i in range(len(all_projects)):
                if i not in project_ranks:
                    project_ranks.append(i)
            
            # Reorder projects
            sorted_projects = [all_projects[idx] for idx in project_ranks]
            
            # Keep top 3 most relevant projects
            customized_data["projects"] = sorted_projects[:3]
            
            logger.info(f"Selected and reordered {len(customized_data['projects'])} most relevant projects")
            
        except Exception as e:
            logger.error(f"Error customizing projects with AI: {e}")
            # Fallback to keeping all projects if AI fails
            if len(all_projects) > 3:
                customized_data["projects"] = all_projects[:3]
    
    def create_custom_resume(self, job_info, output_filename=None):
        """Create a customized resume document based on job requirements."""
        # Extract requirements from job description using AI
        job_requirements = self.extract_job_requirements(job_info["description"])
        
        # Create customized resume data
        customized_data = self.customize_resume_data(job_requirements, job_info)
        
        # Create a new blank document (not using template to avoid duplicate headers)
        doc = Document()
        
        # Set narrow margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)
        
        # Add document structure
        self._add_header(doc, customized_data)
        self._add_professional_summary(doc, customized_data.get("professional_summary", ""))
        self._add_skills_section(doc, customized_data["skills"])
        self._add_work_experience_section(doc, customized_data["work_experience"])
        self._add_education_section(doc, customized_data["education"])
        self._add_projects_section(doc, customized_data.get("projects", []))
        
        # Set output filename if not provided
        if not output_filename:
            company_slug = job_info['company'].replace(" ", "_").replace(".", "")
            output_filename = f"{OUTPUT_DIR}/Resume_{customized_data['personal_info']['last_name']}_{company_slug}_{job_info['job_id']}_{datetime.now().strftime('%Y%m%d')}.docx"
        
        # Save the document
        doc.save(output_filename)
        logger.info(f"Saved customized resume to {output_filename}")
        
        return output_filename
    
    def _add_header(self, doc, customized_data):
        """Add personal information header to the resume with large name."""
        from modules.resume.resume_style_config import FONT_SIZES, FONT_STYLES, SPACING
        
        personal_info = customized_data["personal_info"]
        
        # Name - larger font size
        name_para = doc.add_paragraph()
        name_run = name_para.add_run(f"{personal_info['first_name']} {personal_info['last_name']}")
        name_run.bold = FONT_STYLES["name_bold"]
        name_run.font.size = Pt(FONT_SIZES["name"])
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Contact info
        contact_para = doc.add_paragraph()
        contact_info = []
        
        if "location" in personal_info:
            location = personal_info["location"]
            if "relocation_note" in customized_data:
                location = f"{location} (Relocating to {customized_data['relocation_note']})"
            contact_info.append(location)
        
        if "phone" in personal_info:
            contact_info.append(personal_info["phone"])
        
        if "email" in personal_info:
            contact_info.append(personal_info["email"])
        
        if "linkedin" in personal_info:
            contact_info.append(personal_info["linkedin"])
        
        contact_run = contact_para.add_run(" | ".join(contact_info))
        contact_run.font.size = Pt(FONT_SIZES["contact_info"])
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add horizontal line after contact info (using underscores instead of border)
        border_para = doc.add_paragraph()
        border_para.paragraph_format.space_before = Pt(SPACING["after_name"])
        border_para.paragraph_format.space_after = Pt(SPACING["after_contact_info"])
        border_run = border_para.add_run("_" * 100)  # Underscores as a visual separator
        border_run.font.size = Pt(8)
        border_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def _add_professional_summary(self, doc, summary_text):
        """Add the professional summary section."""
        # Add section heading
        section_heading = doc.add_paragraph()
        heading_run = section_heading.add_run("PROFESSIONAL SUMMARY")
        heading_run.bold = True
        heading_run.font.size = Pt(12)
        section_heading.paragraph_format.space_after = Pt(4)
        
        # Add summary text
        summary_para = doc.add_paragraph()
        summary_run = summary_para.add_run(summary_text)
        summary_run.font.size = Pt(10)
        summary_para.paragraph_format.space_after = Pt(8)
    
    def _add_skills_section(self, doc, skills):
        """Add skills section with emphasis on matched skills."""
        # Add section heading
        section_heading = doc.add_paragraph()
        heading_run = section_heading.add_run("SKILLS")
        heading_run.bold = True
        heading_run.font.size = Pt(12)
        section_heading.paragraph_format.space_after = Pt(4)
        
        # Format and add skills
        skills_para = doc.add_paragraph()
        skills_text = ", ".join(skills)
        skills_run = skills_para.add_run(skills_text)
        skills_run.font.size = Pt(10)
        skills_para.paragraph_format.space_after = Pt(8)
    
    def _add_work_experience_section(self, doc, experiences):
        """Add work experience section with dates right-aligned."""
        # Add section heading
        section_heading = doc.add_paragraph()
        heading_run = section_heading.add_run("WORK EXPERIENCE")
        heading_run.bold = True
        heading_run.font.size = Pt(12)
        section_heading.paragraph_format.space_after = Pt(4)
        
        # Add each experience with table layout for right-aligned dates
        for exp in experiences:
            # Create a table for company/title and dates
            table = doc.add_table(rows=2, cols=2)
            table.autofit = True
            table.allow_autofit = True
            
            # Remove default table borders
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(10)
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            
            # Company name (left cell, top row)
            company_cell = table.cell(0, 0)
            company_para = company_cell.paragraphs[0]
            company_run = company_para.add_run(f"{exp['company']}")
            company_run.bold = True
            company_run.font.size = Pt(11)
            
            if "location" in exp:
                company_para.add_run(f" {exp['location']}")
            
            # Dates (right cell, top row)
            dates_cell = table.cell(0, 1)
            dates_para = dates_cell.paragraphs[0]
            dates_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            if "dates" in exp:
                dates_para.add_run(exp['dates']).font.size = Pt(10)
            
            # Title (left cell, bottom row)
            title_cell = table.cell(1, 0)
            title_para = title_cell.paragraphs[0]
            title_run = title_para.add_run(f"{exp['title']}")
            title_run.italic = True
            title_run.font.size = Pt(10)
            
            # Clear right cell, bottom row
            table.cell(1, 1).text = ""
            
            # Set table to use entire width
            table.autofit = False
            table.width = doc.sections[0].page_width - doc.sections[0].left_margin - doc.sections[0].right_margin
            
            # Bullet points
            if "bullet_points" in exp and isinstance(exp["bullet_points"], list):
                for bullet in exp["bullet_points"]:
                    bullet_para = doc.add_paragraph()
                    bullet_para.style = "List Bullet"
                    bullet_para.add_run(bullet).font.size = Pt(10)
            elif "description" in exp:
                # For descriptions without bullet points
                desc_para = doc.add_paragraph()
                desc_para.add_run(exp["description"]).font.size = Pt(10)
            
            # Add spacing after experience
            doc.add_paragraph().paragraph_format.space_after = Pt(8)
    
    def _add_education_section(self, doc, education):
        """Add education section with dates right-aligned."""
        # Add section heading
        section_heading = doc.add_paragraph()
        heading_run = section_heading.add_run("EDUCATION")
        heading_run.bold = True
        heading_run.font.size = Pt(12)
        section_heading.paragraph_format.space_after = Pt(4)
        
        # Add each education entry with table layout for right-aligned dates
        for edu in education:
            # Create a table for school and dates
            table = doc.add_table(rows=2, cols=2)
            table.autofit = True
            table.allow_autofit = True
            
            # Remove default table borders
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(10)
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            
            # School name (left cell, top row)
            school_cell = table.cell(0, 0)
            school_para = school_cell.paragraphs[0]
            school_run = school_para.add_run(f"{edu['school']}")
            school_run.bold = True
            school_run.font.size = Pt(11)
            
            if "location" in edu:
                school_para.add_run(f" {edu['location']}")
            
            # Dates (right cell, top row)
            dates_cell = table.cell(0, 1)
            dates_para = dates_cell.paragraphs[0]
            dates_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            if "dates" in edu:
                dates_para.add_run(edu['dates']).font.size = Pt(10)
            
            # Degree (left cell, bottom row)
            degree_cell = table.cell(1, 0)
            degree_para = degree_cell.paragraphs[0]
            degree_text = f"{edu.get('degree', '')} {edu.get('field', '')}"
            degree_run = degree_para.add_run(degree_text.strip())
            degree_run.italic = True
            degree_run.font.size = Pt(10)
            
            # Clear right cell, bottom row
            table.cell(1, 1).text = ""
            
            # Set table to use entire width
            table.autofit = False
            table.width = doc.sections[0].page_width - doc.sections[0].left_margin - doc.sections[0].right_margin
            
            # Additional education info
            if "gpa" in edu:
                gpa_para = doc.add_paragraph()
                gpa_para.add_run(f"GPA: {edu['gpa']}").font.size = Pt(10)
            
            if "honors" in edu:
                honors_para = doc.add_paragraph()
                honors_para.add_run(f"Honors: {edu['honors']}").font.size = Pt(10)
            
            # Add spacing after education
            doc.add_paragraph().paragraph_format.space_after = Pt(8)
    
    def _add_projects_section(self, doc, projects):
        """Add projects section with dates right-aligned."""
        # Skip if no projects
        if not projects:
            return
        
        # Add section heading
        section_heading = doc.add_paragraph()
        heading_run = section_heading.add_run("PROJECTS")
        heading_run.bold = True
        heading_run.font.size = Pt(12)
        section_heading.paragraph_format.space_after = Pt(4)
        
        # Add each project with table layout for right-aligned dates
        for project in projects:
            # Create a table for project title and dates
            table = doc.add_table(rows=1, cols=2)
            table.autofit = True
            table.allow_autofit = True
            
            # Remove default table borders
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(10)
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            
            # Project title (left cell)
            title_cell = table.cell(0, 0)
            title_para = title_cell.paragraphs[0]
            title_run = title_para.add_run(f"{project['title']}")
            title_run.bold = True
            title_run.font.size = Pt(11)
            
            # Dates (right cell)
            dates_cell = table.cell(0, 1)
            dates_para = dates_cell.paragraphs[0]
            dates_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            if "dates" in project:
                dates_para.add_run(project['dates']).font.size = Pt(10)
            
            # Set table to use entire width
            table.autofit = False
            table.width = doc.sections[0].page_width - doc.sections[0].left_margin - doc.sections[0].right_margin
            
            # Description
            if "description" in project:
                desc_para = doc.add_paragraph()
                desc_para.add_run(project["description"]).font.size = Pt(10)
            
            # Technologies used
            if "technologies" in project:
                tech_para = doc.add_paragraph()
                tech_run = tech_para.add_run(f"Technologies: {project['technologies']}")
                tech_run.italic = True
                tech_run.font.size = Pt(10)
            
            # Bullet points if available
            if "bullet_points" in project and isinstance(project["bullet_points"], list):
                for bullet in project["bullet_points"]:
                    bullet_para = doc.add_paragraph()
                    bullet_para.style = "List Bullet"
                    bullet_para.add_run(bullet).font.size = Pt(10)
            
            # Add spacing after project
            doc.add_paragraph().paragraph_format.space_after = Pt(8)

    def convert_to_pdf(self, docx_file):
        """Convert the DOCX resume to PDF format."""
        try:
            pdf_file = docx_file.replace('.docx', '.pdf')
            convert(docx_file, pdf_file)
            logger.info(f"Converted resume to PDF: {pdf_file}")
            return pdf_file
        except Exception as e:
            logger.error(f"Error converting to PDF: {e}")
            return None
    
    def cleanup(self):
        """Clean up resources."""
        if self.ai_client:
            ai_close_openai_client(self.ai_client)
            logger.info("Closed OpenAI client")


def parse_arguments():
    """Parse command line arguments."""
    parser = argparse.ArgumentParser(description='Generate AI-powered customized resumes based on job descriptions.')
    parser.add_argument('--job_id', required=True, help='Job ID from the applications CSV')
    parser.add_argument('--output', help='Output filename for the generated resume')
    parser.add_argument('--pdf', action='store_true', help='Convert the resume to PDF format')
    parser.add_argument('--log', choices=['DEBUG', 'INFO', 'WARNING', 'ERROR'], 
                       default='INFO', help='Set logging level')
    return parser.parse_args()


def main():
    """Main function to run the resume customizer."""
    args = parse_arguments()
    
    # Set logging level
    logger.setLevel(getattr(logging, args.log))
    
    customizer = None
    try:
        # Initialize the resume customizer
        customizer = AIResumeCustomizer()
        
        # Load job description
        job_info = customizer.load_job_description(int(args.job_id))
        
        # Create custom resume
        output_file = args.output if args.output else None
        docx_file = customizer.create_custom_resume(job_info, output_file)
        
        # Convert to PDF if requested
        if args.pdf:
            pdf_file = customizer.convert_to_pdf(docx_file)
            if pdf_file:
                logger.info(f"Resume customized and saved as PDF: {pdf_file}")
        else:
            logger.info(f"Resume customized and saved: {docx_file}")
        
    except Exception as e:
        logger.error(f"Error generating resume: {e}")
        raise
    finally:
        # Clean up resources
        if customizer:
            customizer.cleanup()


if __name__ == "__main__":
    main()
