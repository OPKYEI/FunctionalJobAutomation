"""
Integration module for the Resume Customizer with LinkedIn job automation system.
This module provides functions to generate customized resumes based on job descriptions.
"""

import os
import sys
import json
import logging
from datetime import datetime
from pathlib import Path

# Configure basic logging for this module
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    datefmt='%Y-%m-%d %H:%M:%S'
)
logger = logging.getLogger(__name__)

# Define paths
CUSTOMIZED_RESUMES_DIR = "all resumes/customized"
RESUME_DATA_PATH = "data/resume_data.json"

def print_lg(msg, end="\n", flush=False, pretty=False):
    """
    Log function replacement for integration with the main application's logging.
    
    Args:
        msg: Message to log
        end: End character for print (default: newline)
        flush: Whether to flush output
        pretty: Whether to format the output for pretty printing
    """
    # Log to the module's logger
    logger.info(msg)
    
    # Try to use the global print_lg if available
    try:
        # Import the main application's logging function
        from modules.helpers import print_lg as global_print_lg
        global_print_lg(msg, end=end, flush=flush, pretty=pretty)
    except ImportError:
        # If we can't import it, use regular print
        print(msg, end=end, flush=flush)
def create_custom_resume(job_id, title, company, work_location, work_style, job_description):
    """
    Create a custom resume based on job details and verify it was created successfully.
    
    Args:
        job_id: Unique identifier for the job
        title: Job title
        company: Company name
        work_location: Job location
        work_style: Remote, hybrid, on-site, etc.
        job_description: Full job description text
        
    Returns:
        str: Absolute path to the created resume if successful, None otherwise
    """
    try:
        # Create directories for custom resumes
        custom_resume_dir = os.path.abspath("all resumes/customized")
        if not os.path.exists(custom_resume_dir):
            os.makedirs(custom_resume_dir)
            print_lg(f"📁 Created custom resume directory: {custom_resume_dir}")
        
        # Clean company and title names for filename
        safe_company = ''.join(c if c.isalnum() else '_' for c in company)
        safe_title = ''.join(c if c.isalnum() else '_' for c in title)
        
        # Create a filename with job ID, title, and company
        timestamp = datetime.now().strftime('%Y%m%d%H%M%S')
        resume_filename = f"Resume_{safe_title}_{safe_company}_{job_id}_{timestamp}.docx"
        custom_resume_path = os.path.abspath(os.path.join(custom_resume_dir, resume_filename))
        
        print_lg(f"🔄 Creating custom resume for {company}: {title}")
        print_lg(f"   Target path: {custom_resume_path}")
        
        # USE AI-POWERED RESUME CUSTOMIZER IF AVAILABLE
        try:
            # Check if we're using free AI mode or regular API
            try:
                from config.settings import use_free_ai
            except ImportError:
                try:
                    from config.secrets import use_free_ai
                except ImportError:
                    # If use_free_ai is not defined in either file, default to True
                    use_free_ai = True
                
            from modules.resume.ai_resume_customizer import AIResumeCustomizer
            from modules.ai.openaiConnections import free_ai_completion, extract_json_from_text
            from g4f.client import Client
            
            # Create job info dictionary for the customizer
            job_info = {
                'job_id': job_id,
                'title': title,
                'company': company,
                'work_location': work_location,
                'work_style': work_style,
                'description': job_description
            }
            
            # Modified approach for using AI resume customizer
            if hasattr(AIResumeCustomizer, '_handle_free_ai_mode'):
                # If the customizer already has the free AI mode handler, use it normally
                customizer = AIResumeCustomizer()
                ai_resume_path = customizer.create_custom_resume(job_info, custom_resume_path)
                customizer.cleanup()
            else:
                # Manual adaptation for free AI mode - create a basic customized resume
                # Create a blank document
                from docx import Document
                from docx.shared import Pt, Inches
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                
                doc = Document()
                
                # Add a header with name and job target
                header = doc.add_heading(f"Resume for {title} at {company}", 0)
                header.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Get customized content using free AI
                try:
                    # Extract key skills
                    skills_prompt = f"""
                    Extract the top 10-15 technical skills and keywords from this job description. 
                    Format as a comma-separated list only.
                    
                    Job Description:
                    {job_description}
                    """
                    
                    # Get skills using free AI
                    client = Client()
                    skills_response = client.chat.completions.create(
                        model="gpt-3.5-turbo",
                        messages=[{"role": "user", "content": skills_prompt}]
                    )
                    skills_list = skills_response.choices[0].message.content.strip()
                    
                    # Add skills section
                    doc.add_heading("Key Skills", 1)
                    doc.add_paragraph(skills_list)
                    
                    # Create a tailored professional summary
                    summary_prompt = f"""
                    Write a concise, personalized professional summary (3 sentences) for a resume
                    targeting this job:
                    
                    Job Title: {title}
                    Company: {company}
                    
                    Job Description:
                    {job_description}
                    
                    Focus on matching the candidate to this specific role.
                    """
                    
                    summary_response = client.chat.completions.create(
                        model="gpt-3.5-turbo",
                        messages=[{"role": "user", "content": summary_prompt}]
                    )
                    summary_text = summary_response.choices[0].message.content.strip()
                    
                    # Add professional summary
                    doc.add_heading("Professional Summary", 1)
                    doc.add_paragraph(summary_text)
                    
                    # This is a simplified implementation - a production version would
                    # add more sections like Experience, Education, etc.
                    
                except Exception as ai_error:
                    print_lg(f"Error generating content with free AI: {ai_error}")
                    # Add a placeholder summary and skills section
                    doc.add_heading("Professional Summary", 1)
                    doc.add_paragraph(f"Experienced professional seeking the {title} position at {company}.")
                    
                    doc.add_heading("Key Skills", 1)
                    doc.add_paragraph("Python, SQL, Data Analysis, Problem Solving, Communication")
                
                # Save the document
                doc.save(custom_resume_path)
                
                # Try to convert to PDF if docx2pdf is available
                try:
                    from docx2pdf import convert
                    pdf_path = custom_resume_path.replace('.docx', '.pdf')
                    convert(custom_resume_path, pdf_path)
                    print_lg(f"Converted resume to PDF: {pdf_path}")
                    ai_resume_path = pdf_path
                except:
                    ai_resume_path = custom_resume_path
            
            # Verify the file exists
            if ai_resume_path and os.path.exists(ai_resume_path):
                print_lg(f"✅ AI-powered resume created successfully: {ai_resume_path}")
                return ai_resume_path
                
        except Exception as ai_error:
            print_lg(f"⚠️ AI resume creation failed, falling back to basic resume: {str(ai_error)}")
            # Continue with fallback method

        # FALLBACK: CREATE BASIC RESUME
        # For now, just copy the default resume as a placeholder
        try:
            import shutil
            from config.questions import default_resume_path
            
            default_resume = os.path.abspath(default_resume_path)
            
            if os.path.exists(default_resume):
                # Determine target extension
                target_ext = os.path.splitext(default_resume)[1]
                fallback_path = custom_resume_path.replace('.docx', target_ext)
                
                shutil.copy(default_resume, fallback_path)
                
                print_lg(f"✅ Basic resume created as fallback: {fallback_path}")
                return fallback_path
            else:
                print_lg(f"❌ Default resume not found at {default_resume}")
                
        except Exception as copy_error:
            print_lg(f"❌ Failed to copy default resume: {str(copy_error)}")
        
        # After resume creation, perform thorough verification
        for check_path in [custom_resume_path, custom_resume_path.replace('.docx', '.pdf')]:
            if os.path.exists(check_path):
                file_size = os.path.getsize(check_path)
                creation_time = datetime.fromtimestamp(os.path.getctime(check_path))
                
                if file_size > 0:
                    print_lg(f"✅ Custom resume created and verified: {os.path.basename(check_path)}")
                    print_lg(f"   Full path: {check_path}")
                    print_lg(f"   File size: {file_size} bytes")
                    print_lg(f"   Created at: {creation_time}")
                    
                    # Return the ABSOLUTE path
                    return check_path
        
        print_lg(f"❌ Error: No valid resume file was created")
        return None
            
    except Exception as e:
        print_lg(f"❌ Error creating custom resume: {str(e)}")
        # Stack trace for debugging
        import traceback
        print_lg(f"   Stack trace: {traceback.format_exc()}")
        return None

def get_resume_path_for_job(job_id, use_default=False):
    """
    Check if a customized resume already exists for this job ID.
    If not found and use_default is True, returns the default resume path.
    
    Args:
        job_id (str): LinkedIn job ID
        use_default (bool): Whether to return default resume path if no custom resume found
        
    Returns:
        str: Path to existing resume if found, default resume if use_default=True, None otherwise
    """
    # Check if a custom resume exists
    if not os.path.exists(CUSTOMIZED_RESUMES_DIR):
        if use_default:
            from config.questions import default_resume_path
            return os.path.abspath(default_resume_path)
        return None
        
    # List all files in the directory
    for filename in os.listdir(CUSTOMIZED_RESUMES_DIR):
        if str(job_id) in filename and filename.endswith(('.docx', '.pdf')):
            return os.path.abspath(os.path.join(CUSTOMIZED_RESUMES_DIR, filename))
    
    # If no custom resume found and use_default is True, return default resume
    if use_default:
        from config.questions import default_resume_path
        return os.path.abspath(default_resume_path)
            
    return None